# import fitz  # PyMuPDF
# from docx import Document
# import re
# import streamlit as st
# import os
# from dotenv import load_dotenv
# from google import genai
# from google.genai import types
# from PIL import Image
# from io import BytesIO

# # Load environment variables
# load_dotenv()

# # Configure Gemini client
# client = genai.Client(api_key=os.getenv("GOOGLE_API_KEY"))

# # Extract text from PDF or TXT
# def extract_text_from_file(uploaded_file):
#     if uploaded_file.type == "application/pdf":
#         doc = fitz.open(stream=uploaded_file.read(), filetype="pdf")
#         return "".join([page.get_text() for page in doc])
#     else:
#         return uploaded_file.read().decode("utf-8")

# # Extract keywords from JD
# def extract_keywords(text, top_n=15):
#     words = re.findall(r'\b\w+\b', text.lower())
#     ignore = set(['and', 'with', 'the', 'for', 'you', 'are', 'that', 'have', 'job', 'your', 'will'])
#     freq = {}
#     for word in words:
#         if len(word) > 3 and word not in ignore:
#             freq[word] = freq.get(word, 0) + 1
#     sorted_keywords = sorted(freq.items(), key=lambda x: x[1], reverse=True)
#     return [word for word, count in sorted_keywords[:top_n]]

# # Compute keyword match
# def keyword_match_score(resume, keywords):
#     resume_words = resume.lower()
#     match_count = sum(1 for word in keywords if word in resume_words)
#     return int((match_count / len(keywords)) * 100)

# # Save as DOCX
# def save_as_docx(text, filename):
#     doc = Document()
#     for line in text.split("\n"):
#         doc.add_paragraph(line)
#     doc.save(filename)

# # Generate interview questions
# def generate_interview_questions(jd_text, resume_text, num_questions=10):
#     prompt = f"""
# You are an expert career coach.

# Based on the following job description and candidate resume, generate {num_questions} likely interview questions that the candidate may face.

# Job Description:
# {jd_text}

# Candidate Resume:
# {resume_text}

# Provide a numbered list of questions only.
# """
#     response = client.models.generate_content(
#         model="gemini-2.5-flash-preview-05-20",
#         contents=prompt
#     )
#     questions_text = response.text.strip()
#     return questions_text.split('\n')

# # ---------- Streamlit UI Setup ----------
# st.set_page_config(page_title="SmartJob AI: Resume Tailoring, Interview Questions & Visual Templates", layout="wide")
# st.title("🧑‍💻 SmartJob AI: Resume Tailoring, Interview Questions & Visual Templates")
# st.markdown("Upload your resume and job description to generate a tailored resume, cover letter, likely interview questions, and even a resume layout image!")

# col1, col2 = st.columns(2)

# # Resume uploader
# with col1:
#     resume_file = st.file_uploader("📄 Upload Resume (PDF or .txt)", type=["pdf", "txt"])

# # JD input section with upload or text area
# with col2:
#     jd_input_mode = st.radio("📋 Job Description Input Mode", ["Upload File", "Type Manually"])
#     jd_text = ""
#     jd_file = None
#     if jd_input_mode == "Upload File":
#         jd_file = st.file_uploader("📎 Upload Job Description (PDF or .txt)", type=["pdf", "txt"])
#         if jd_file:
#             jd_text = extract_text_from_file(jd_file)
#     else:
#         jd_text = st.text_area("🖊️ Paste or Type the Job Description Here", height=250)

# # Tone selection
# tone = st.selectbox("✍️ Select Writing Tone", ["Formal", "Conversational", "Creative"])

# # Image prompt input
# custom_image_prompt = st.text_area("🎨 Enter Visual Prompt for Resume Design (style, layout, colors, etc.):", height=100)

# # Generate All Outputs
# if st.button("🚀 Generate All Outputs") and resume_file and jd_text.strip():
#     resume_text = extract_text_from_file(resume_file)

#     keywords = extract_keywords(jd_text)
#     match_score = keyword_match_score(resume_text, keywords)

#     st.subheader("📊 ATS Keyword Match Score")
#     st.progress(match_score / 100)
#     st.markdown(f"**{match_score}%** match with the job description keywords.")

#     with st.spinner("Generating tailored resume, cover letter, interview questions, and image..."):
#         try:
#             # Resume and Cover Letter
#             prompt = f"""
# You are a resume and cover letter writing expert.

# Task:
# - Rewrite the resume to closely match the job description using a {tone.lower()} tone.
# - Generate a professional, tailored cover letter for the position.

# Job Description:
# {jd_text}

# Original Resume:
# {resume_text}

# Output the tailored resume first, then the cover letter. Use markdown headings for separation.
#             """
#             response = client.models.generate_content(
#                 model="gemini-2.5-flash-preview-05-20",
#                 contents=prompt
#             )
#             output_text = response.text
#             st.session_state.tailored_text = output_text

#             st.subheader("📄 Tailored Resume & Cover Letter")
#             editable = st.text_area("Review and edit the generated content if needed:", value=output_text, height=500)
#             st.download_button("⬇️ Download as TXT", editable, file_name="tailored_documents.txt")
#             save_as_docx(editable, "tailored_documents.docx")
#             with open("tailored_documents.docx", "rb") as f:
#                 st.download_button("⬇️ Download as DOCX", f, file_name="tailored_documents.docx")

#             # Interview Questions
#             questions = generate_interview_questions(jd_text, resume_text, num_questions=10)
#             st.subheader("🎯 Likely Interview Questions")
#             for q in questions:
#                 st.markdown(q)

#             # Image Generation
#             if custom_image_prompt.strip():
#                 image_prompt_to_use = custom_image_prompt.strip()
#                 st.subheader("🖼️ AI-Generated Resume Layout Image")
#                 st.markdown(f"🧠 Using image prompt: _{image_prompt_to_use}_")

#                 try:
#                     image_response = client.models.generate_content(
#                         model="gemini-2.0-flash-preview-image-generation",
#                         contents=image_prompt_to_use,
#                         config=types.GenerateContentConfig(
#                             response_modalities=['TEXT', 'IMAGE']
#                         )
#                     )
#                     for part in image_response.candidates[0].content.parts:
#                         if part.text is not None:
#                             st.markdown(part.text)
#                         elif part.inline_data is not None:
#                             image = Image.open(BytesIO(part.inline_data.data))
#                             st.image(image, caption="🧠 AI-Generated Resume Design (Gemini)")
#                             image_path = "gemini_generated_image.png"
#                             image.save(image_path)
#                             with open(image_path, "rb") as img_file:
#                                 st.download_button("⬇️ Download Image", img_file, file_name="resume_design.png")
#                 except Exception as e:
#                     st.error(f" Failed to generate image: {e}")
#             else:
#                 st.info("ℹ️ No image prompt entered, skipping image generation.")

#         except Exception as e:
#             st.error(f"Error generating content: {e}")


import fitz  # PyMuPDF
from docx import Document
import re
import streamlit as st
import os
from dotenv import load_dotenv
from google import genai
from google.genai import types
from PIL import Image
from io import BytesIO
from collections import Counter
import math

# Load environment variables
load_dotenv()

# Configure Gemini client
client = genai.Client(api_key=os.getenv("GOOGLE_API_KEY"))

# Extract text from PDF or TXT
def extract_text_from_file(uploaded_file):
    if uploaded_file.type == "application/pdf":
        doc = fitz.open(stream=uploaded_file.read(), filetype="pdf")
        return "".join([page.get_text() for page in doc])
    else:
        return uploaded_file.read().decode("utf-8")

# TF-IDF based keyword extraction
def extract_keywords_tfidf(text, top_n=20):
    words = re.findall(r'\b\w+\b', text.lower())
    ignore = set(['and', 'with', 'the', 'for', 'you', 'are', 'that', 'have', 'job', 'your', 'will', 'this', 'from', 'but', 'not', 'all', 'any', 'can', 'our', 'has'])
    words_filtered = [w for w in words if len(w) > 3 and w not in ignore]

    tf = Counter(words_filtered)
    total_terms = len(words_filtered)
    tf = {word: count/total_terms for word, count in tf.items()}

    # For IDF, use a simple heuristic: words appearing many times are less informative
    idf = {}
    for word in tf:
        # Simple IDF approximation (the rarer the word, the higher the score)
        idf[word] = 1 / (math.log(1 + tf[word]*total_terms) + 1e-6)
    
    tfidf = {word: tf[word]*idf[word] for word in tf}
    sorted_keywords = sorted(tfidf.items(), key=lambda x: x[1], reverse=True)
    return [word for word, score in sorted_keywords[:top_n]]

# Compute keyword match score with detailed metrics
def keyword_match_score(resume_text, keywords):
    resume_words = re.findall(r'\b\w+\b', resume_text.lower())
    resume_word_set = set(resume_words)
    matched = [k for k in keywords if k in resume_word_set]
    match_count = len(matched)
    match_percent = (match_count / len(keywords)) * 100 if keywords else 0

    # Keyword density in resume (for matched keywords only)
    freq = Counter(resume_words)
    density = {k: freq[k]/len(resume_words) for k in matched}

    missing = [k for k in keywords if k not in resume_word_set]

    return {
        "match_percent": int(match_percent),
        "matched_keywords": matched,
        "missing_keywords": missing,
        "density": density,
    }

# Save as DOCX
def save_as_docx(text, filename):
    doc = Document()
    for line in text.split("\n"):
        doc.add_paragraph(line)
    doc.save(filename)

# Generate interview questions
def generate_interview_questions(jd_text, resume_text, num_questions=10):
    prompt = f"""
You are an expert career coach.

Based on the following job description and candidate resume, generate {num_questions} likely interview questions that the candidate may face.

Job Description:
{jd_text}

Candidate Resume:
{resume_text}

Provide a numbered list of questions only.
"""
    response = client.models.generate_content(
        model="gemini-2.5-flash-preview-05-20",
        contents=prompt
    )
    questions_text = response.text.strip()
    return questions_text.split('\n')

# Generate tailored resume and cover letter
def generate_tailored_resume_coverletter(jd_text, resume_text, tone):
    prompt = f"""
You are a resume and cover letter writing expert.

Task:
- Rewrite the resume to closely match the job description using a {tone.lower()} tone.
- Generate a professional, tailored cover letter for the position.

Job Description:
{jd_text}

Original Resume:
{resume_text}

Output the tailored resume first, then the cover letter. Use markdown headings for separation.
"""
    response = client.models.generate_content(
        model="gemini-2.5-flash-preview-05-20",
        contents=prompt
    )
    return response.text

# Generate summary of resume
def generate_resume_summary(resume_text):
    prompt = f"""
Summarize the key skills and experience of the following resume in 3-4 concise bullet points:

Resume:
{resume_text}
"""
    response = client.models.generate_content(
        model="gemini-2.5-flash-preview-05-20",
        contents=prompt
    )
    return response.text.strip()

# Extract email and phone from resume text
def extract_email_phone(text):
    email_match = re.search(r'[\w\.-]+@[\w\.-]+\.\w+', text)
    phone_match = re.search(r'(\+?\d{1,3}[-.\s]?)?(\(?\d{3}\)?[-.\s]?){1,2}\d{4}', text)
    email = email_match.group(0) if email_match else None
    phone = phone_match.group(0) if phone_match else None
    return email, phone

# ---------- Streamlit UI Setup ----------
st.set_page_config(page_title="SmartJob AI: Resume Tailoring, Interview Questions & Visual Templates", layout="wide")
st.title("🧑‍💻 SmartJob AI: Resume Tailoring, Interview Questions & Visual Templates")
st.markdown("Upload your resume and job description to generate a tailored resume, cover letter, interview questions, and AI-generated resume layout image!")

col1, col2 = st.columns(2)

# Resume uploader
with col1:
    resume_file = st.file_uploader("📄 Upload Resume (PDF or .txt)", type=["pdf", "txt"])

# JD input section with upload or text area
with col2:
    jd_input_mode = st.radio("📋 Job Description Input Mode", ["Upload File", "Type Manually"])
    jd_text = ""
    jd_file = None
    if jd_input_mode == "Upload File":
        jd_file = st.file_uploader("📎 Upload Job Description (PDF or .txt)", type=["pdf", "txt"])
        if jd_file:
            jd_text = extract_text_from_file(jd_file)
    else:
        jd_text = st.text_area("🖊️ Paste or Type the Job Description Here", height=250)

# Tone selection for writing style
tone = st.selectbox("✍️ Select Writing Tone", ["Formal", "Conversational", "Creative"])

# Image prompt input
custom_image_prompt = st.text_area("🎨 Enter Visual Prompt for Resume Design (style, layout, colors, etc.):", height=100)

# Generate All Outputs
if st.button("🚀 Generate All Outputs") and resume_file and jd_text.strip():
    with st.spinner("Processing files..."):
        try:
            resume_text = extract_text_from_file(resume_file)
            # Extract email & phone
            email, phone = extract_email_phone(resume_text)

            # Keyword extraction & ATS metrics
            keywords = extract_keywords_tfidf(jd_text)
            ats_metrics = keyword_match_score(resume_text, keywords)

            st.subheader("📊 ATS Keyword Match Score")
            st.progress(ats_metrics["match_percent"] / 100)
            st.markdown(f"**{ats_metrics['match_percent']}%** match with job description keywords.")
            st.markdown(f"**Matched Keywords:** {', '.join(ats_metrics['matched_keywords'])}")
            if ats_metrics['missing_keywords']:
                st.markdown(f"**Missing Keywords:** {', '.join(ats_metrics['missing_keywords'])}")

            # Show extracted contact info (if found)
            if email or phone:
                st.subheader("📞 Extracted Contact Info from Resume")
                if email:
                    st.markdown(f"📧 Email: {email}")
                if phone:
                    st.markdown(f"📱 Phone: {phone}")

            # Resume summary
            st.subheader("📝 Resume Summary")
            summary = generate_resume_summary(resume_text)
            st.markdown(summary)

            # Generate tailored resume & cover letter
            st.subheader("📄 Generating Tailored Resume & Cover Letter...")
            tailored_text = generate_tailored_resume_coverletter(jd_text, resume_text, tone)
            st.session_state.tailored_text = tailored_text
            st.text_area("Review and edit the generated content if needed:", value=tailored_text, height=500)

            # Separate download buttons for resume & cover letter
            if "## Tailored Resume" in tailored_text and "## Cover Letter" in tailored_text:
                resume_part = tailored_text.split("## Cover Letter")[0].strip()
                cover_letter_part = "## Cover Letter" + tailored_text.split("## Cover Letter")[1]

                st.download_button("⬇️ Download Tailored Resume as TXT", resume_part, file_name="tailored_resume.txt")
                st.download_button("⬇️ Download Cover Letter as TXT", cover_letter_part, file_name="cover_letter.txt")

                save_as_docx(tailored_text, "tailored_documents.docx")
                with open("tailored_documents.docx", "rb") as f:
                    st.download_button("⬇️ Download Combined Resume & Cover Letter as DOCX", f, file_name="tailored_documents.docx")
            else:
                # Fallback download button
                st.download_button("⬇️ Download Generated Document as TXT", tailored_text, file_name="tailored_documents.txt")

            # Generate interview questions
            st.subheader("🎯 Likely Interview Questions")
            questions = generate_interview_questions(jd_text, resume_text, num_questions=10)
            for q in questions:
                st.markdown(q)

            # Download interview questions text
            st.download_button("⬇️ Download Interview Questions", "\n".join(questions), file_name="interview_questions.txt")

            # Image Generation
            if custom_image_prompt.strip():
                image_prompt_to_use = custom_image_prompt.strip()
                st.subheader("🖼️ AI-Generated Resume Layout Image")
                st.markdown(f"🧠 Using image prompt: _{image_prompt_to_use}_")

                try:
                    image_response = client.models.generate_content(
                        model="gemini-2.0-flash-preview-image-generation",
                        contents=image_prompt_to_use,
                        config=types.GenerateContentConfig(
                            response_modalities=['TEXT', 'IMAGE']
                        )
                    )
                    for part in image_response.candidates[0].content.parts:
                        if part.text is not None:
                            st.markdown(part.text)
                        elif part.inline_data is not None:
                            image = Image.open(BytesIO(part.inline_data.data))
                            st.image(image, caption="🧠 AI-Generated Resume Design (Gemini)")
                            image_path = "gemini_generated_image.png"
                            image.save(image_path)
                            with open(image_path, "rb") as img_file:
                                st.download_button("⬇️ Download Image", img_file, file_name="resume_design.png")
                except Exception as e:
                    st.error(f"Failed to generate image: {e}")
            else:
                st.info("ℹ️ No image prompt entered, skipping image generation.")

        except Exception as e:
            st.error(f"Error processing files or generating content: {e}")

else:
    st.info("Please upload both resume and job description to get started.")

